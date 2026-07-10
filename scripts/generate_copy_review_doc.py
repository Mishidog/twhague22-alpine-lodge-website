import json
import os
import subprocess
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOCX = ROOT / "Alpine_Lodge_Website_Copy_Review.docx"
OUTPUT_MD = ROOT / "Alpine_Lodge_Website_Copy_Review.md"


def load_site_data():
    node = os.environ.get("NODE_BIN", "node")
    script = """
      import * as data from './data/site.js';
      import * as catalog from './data/knowledgeCatalog.mjs';
      const keys = [
        'site',
        'nav',
        'highlights',
        'rateSignal',
        'guestQuotes',
        'rooms',
        'amenities',
        'amenityHighlights',
        'distances',
        'experiences',
        'tripGuides',
        'faqs',
        'homepageFaqs',
        'partnerLinks'
      ];
      const out = {};
      for (const key of keys) out[key] = data[key];
      for (const key of ['catalogUpdated', 'lodgingServices', 'rateKnowledge', 'policyKnowledge']) {
        out[key] = catalog[key];
      }
      console.log(JSON.stringify(out, null, 2));
    """
    result = subprocess.run(
        [node, "--no-warnings", "--input-type=module", "-e", script],
        cwd=ROOT,
        check=True,
        text=True,
        stdout=subprocess.PIPE,
    )
    return json.loads(result.stdout)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    for paragraph in cell.paragraphs:
      paragraph.paragraph_format.space_after = Pt(0)


def style_doc(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 18, "000000"),
        ("Heading 2", 14, "000000"),
        ("Heading 3", 12, "434343"),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = False
        style.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        style.paragraph_format.space_after = Pt(5)


def add_title(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Alpine Lodge Website Copy Review")
    run.font.name = "Arial"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.add_run(f"Prepared for client copy review | Generated {date.today().isoformat()}")
    subtitle.runs[0].font.color.rgb = RGBColor(85, 85, 85)

    doc.add_paragraph(
        "This document organizes the current website copy by where it appears on the site. "
        "It includes global navigation/footer copy, page copy, reusable button text, FAQ copy, "
        "experience page copy, trip guide copy, image alt text, and key SEO metadata for review."
    )


def add_table(doc, rows):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.85)
    table.columns[1].width = Inches(4.65)
    header = table.rows[0].cells
    set_cell_text(header[0], "Element", True)
    set_cell_text(header[1], "Copy", True)
    set_cell_shading(header[0], "F2F4F7")
    set_cell_shading(header[1], "F2F4F7")
    for label, copy in rows:
        row = table.add_row().cells
        set_cell_text(row[0], label)
        set_cell_text(row[1], copy)
    doc.add_paragraph()


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(str(item), style="List Number")


def add_page(doc, title, metadata, sections):
    doc.add_heading(title, level=1)
    if metadata:
        add_table(doc, [("SEO title", metadata.get("title", "")), ("SEO description", metadata.get("description", ""))])
    for section_title, rows in sections:
        doc.add_heading(section_title, level=2)
        add_table(doc, rows)


def make_markdown_line(label, value):
    return f"- **{label}:** {value}\n"


def append_md_page(lines, title, metadata, sections):
    lines.append(f"\n# {title}\n\n")
    if metadata:
        lines.append("## SEO Metadata\n\n")
        lines.append(make_markdown_line("SEO title", metadata.get("title", "")))
        lines.append(make_markdown_line("SEO description", metadata.get("description", "")))
    for section_title, rows in sections:
        lines.append(f"\n## {section_title}\n\n")
        for label, copy in rows:
            lines.append(make_markdown_line(label, copy))


def build_document(data):
    doc = Document()
    style_doc(doc)
    add_title(doc)

    md = [
        "# Alpine Lodge Website Copy Review\n\n",
        f"Prepared for client copy review | Generated {date.today().isoformat()}\n\n",
        "This document organizes the current website copy by where it appears on the site.\n",
    ]

    site = data["site"]
    nav = data["nav"]
    rate = data["rateSignal"]

    global_sections = [
        (
            "Header",
            [
                ("Brand", site["name"]),
                ("Top strip address", "850 Sawmill Lane, Davis, WV"),
                ("Top strip phone", site["phone"]),
                ("Navigation", ", ".join(item["label"] for item in nav)),
                ("Rate badge", rate["eyebrow"]),
                ("Rate badge note", rate["headerNote"]),
                ("Primary CTA", "Reserve Now"),
            ],
        ),
        (
            "Footer",
            [
                ("Footer eyebrow", "Ready for Davis?"),
                ("Footer headline", "Book the Davis room. Spend the trip outside."),
                ("Footer body", "Use Alpine Lodge for Blackwater Falls, Canaan Valley, Timberline, Thomas, and nearby Tucker County stops."),
                ("Footer brand heading", "Alpine Lodge"),
                ("Footer brand body", "Rooms from $99 on select dates in Davis, West Virginia for travelers who came for waterfalls, trails, skiing, food, music, and mountain air."),
                ("Footer CTA", "Check Availability"),
                ("Footer site links", ", ".join(item["label"] for item in nav) + ", Crew Lodging, Contact, Photo Credits"),
                ("Footer social labels", ", ".join(item["label"] for item in site["socials"])),
                ("Footer note", "Experience links are provided for planning. Confirm hours, prices, conditions, and policies with each operator."),
                ("Footer bottom", "Alpine Lodge. Davis, West Virginia lodging near Blackwater Falls."),
            ],
        ),
        (
            "Reusable Buttons And Links",
            [
                ("Primary default", "Reserve Now"),
                ("Secondary default", "Explore"),
                ("Booking button icon meaning", "Opens the direct Cloudbeds reservation site in a new tab."),
                ("Compact card CTA", "View details"),
                ("Full card CTA", "Plan this trip"),
            ],
        ),
    ]
    add_page(doc, "Global Copy", None, global_sections)
    append_md_page(md, "Global Copy", None, global_sections)

    home_sections = [
        (
            "Hero",
            [
                ("Eyebrow", "Rooms from $99 on select dates in Davis, West Virginia"),
                ("Headline", "Your Davis basecamp for Blackwater Falls, Canaan Valley, Timberline & Thomas."),
                ("Body", "Check in on Sawmill Lane, then point the day toward waterfalls, trails, skiing, food, music, and mountain air."),
                ("Primary CTA", "Reserve Your Room"),
                ("Secondary CTA", "Explore The Area"),
                ("Highlight facts", "; ".join(data["highlights"][:4])),
                ("Image alt text", "Alpine Lodge exterior in Davis, West Virginia"),
            ],
        ),
        (
            "Quick Plan Icons",
            [
                ("Activity links", "Waterfalls, Skiing, Hiking, Biking, Food & music, Pet-friendly rooms"),
            ],
        ),
        (
            "Positioning Section",
            [
                ("Eyebrow", "The Alpine Lodge lane"),
                ("Headline", "Not a luxury resort. A smarter place to start the trip."),
                ("Paragraph 1", "You are booking a room in Davis, not a resort itinerary. That is the point. Alpine Lodge puts value-minded travelers close to the Tucker County places they already searched for."),
                ("Paragraph 2", "Sleep here, keep the plan flexible, and use the rest of the trip spend for lift tickets, local food, gear, music, gas, and one more stop before heading home."),
            ],
        ),
        (
            "Amenities Intro",
            [
                ("Eyebrow", "Essentials at a glance"),
                ("Headline", "Check the room basics in ten seconds."),
                ("Body", "Private bath, Wi-Fi, refrigerator, TV, direct booking, and designated pet-friendly rooms when the right room is available."),
            ],
        ),
    ]
    for amenity in data["amenityHighlights"]:
        home_sections.append((f"Amenity Card: {amenity['title']}", [("Title", amenity["title"]), ("Detail", amenity["detail"])]))
    home_sections.extend(
        [
            (
                "Crew Lodging Callout",
                [
                    ("Eyebrow", "Working in the Davis area?"),
                    ("Headline", "A practical room between long workdays."),
                    ("Body", "Contractors, crews, and traveling workers can use Alpine Lodge for Wi-Fi, in-room refrigerators, laundry access, ice, and group booking options by phone."),
                    ("CTA", "View Crew Lodging"),
                ],
            ),
            (
                "Experience Grid Intro",
                [
                    ("Eyebrow", "Explore from Alpine Lodge"),
                    ("Headline", "One room. A whole weekend of options."),
                    ("Body", "Pick the trip first: waterfall morning, ski day, bike ride, Thomas show, rainy-day food loop, dog-friendly drive, or a little of everything."),
                    ("CTA", "See All Experiences"),
                ],
            ),
            (
                "Pet-Friendly Callout",
                [
                    ("Eyebrow", "Pet-friendly rooms"),
                    ("Headline", "Bring the dog, but book the right room."),
                    ("Body", "Designated pet-friendly rooms are available for mountain weekends and road trips through Davis. Because policies and room availability can change, confirm the current pet details before you arrive."),
                    ("CTA", "Check Pet-Friendly Availability"),
                ],
            ),
            (
                "Location Panel",
                [
                    ("Eyebrow", "Central Davis location"),
                    ("Headline", "About 5 to 20 minutes from the area's big reasons to visit."),
                    ("Body", "Davis works as the center point: Blackwater Falls for the classic view, Canaan Valley and Timberline for mountain recreation, Thomas for music and food, and Alpine Lodge for the room."),
                    ("CTA 1", "View Location Details"),
                    ("CTA 2", "Book The Basecamp"),
                ],
            ),
            (
                "Hub Map",
                [
                    ("Eyebrow", "The center-point advantage"),
                    ("Headline", "Alpine Lodge sits in the middle of the trip, not off to the side."),
                    ("Body", "Davis works because the weekend can go in several directions without changing hotels: waterfall morning, valley afternoon, dinner in town, or a ski day when the snow is right."),
                    ("Map nodes", "Alpine Lodge Davis, WV; Blackwater Falls ~5 min; Thomas ~10 min; Canaan Valley ~15 min; Timberline ~20 min; White Grass ~20 min; Dolly Sods Plan ahead"),
                ],
            ),
            (
                "Rooms Strip",
                [
                    ("Eyebrow", "Rooms for after the day outside"),
                    ("Headline", "Shower, fridge, Wi-Fi, TV, sleep. Then go again."),
                    ("Body", "The room is not the destination. It is where you drop the bags, cool the drinks, charge the phone, and reset for tomorrow."),
                    ("CTA", "Check Availability"),
                ],
            ),
            (
                "Rate Panel",
                [
                    ("Eyebrow", rate["eyebrow"]),
                    ("Headline", rate["headline"]),
                    ("Body", rate["body"]),
                    ("CTA 1", "How Rates Work"),
                    ("CTA 2", rate["cta"]),
                ],
            ),
            (
                "Trip Guides Intro",
                [
                    ("Eyebrow", "Practical trip guides"),
                    ("Headline", "Plans visitors can use before they call."),
                    ("Body", "These guides are written for the likely Alpine guest: someone who wants the mountains, a fair room, and a weekend worth the drive."),
                ],
            ),
            (
                "Selected 5/5 Google Reviews",
                [
                    ("Eyebrow", "Selected 5/5 Google reviews"),
                    ("Headline", "Clean rooms, helpful service, and a practical stay."),
                    ("Body", "These verified guest excerpts speak to the things Alpine Lodge visitors tend to care about most."),
                ] + [
                    (quote["name"], f"{quote['rating']}/5 — {quote['text']} Context: {quote['context']}")
                    for quote in data["guestQuotes"]
                ],
            ),
            (
                "FAQ Teaser",
                [
                    ("Eyebrow", "Good answers build confidence"),
                    ("Headline", "Know what to expect before you arrive."),
                    ("Body", "Clear details about pets, check-in, Wi-Fi, ski access, local food, and booking help travelers decide without hunting across three tabs."),
                    ("CTA", "Read all FAQs"),
                    ("Questions shown", "; ".join(faq["question"] for faq in data["homepageFaqs"])),
                ],
            ),
            (
                "Final Booking",
                [
                    ("Headline", "Make Davis your launch point."),
                    ("Body", "Reserve Alpine Lodge, then build the trip around waterfalls, ski areas, trails, music, food, and the people you came with."),
                    ("CTA", "Reserve Now"),
                ],
            ),
        ]
    )
    add_page(
        doc,
        "Homepage",
        {
            "title": "Davis WV Lodging Near Blackwater Falls & Canaan Valley",
            "description": "Alpine Lodge offers Davis, WV rooms from $99 on select dates near Blackwater Falls, Canaan Valley, Timberline Mountain, White Grass, Thomas, trails, skiing, food, and music.",
        },
        home_sections,
    )
    append_md_page(
        md,
        "Homepage",
        {
            "title": "Davis WV Lodging Near Blackwater Falls & Canaan Valley",
            "description": "Alpine Lodge offers Davis, WV rooms from $99 on select dates near Blackwater Falls, Canaan Valley, Timberline Mountain, White Grass, Thomas, trails, skiing, food, and music.",
        },
        home_sections,
    )

    rooms_sections = [
        (
            "Hero",
            [
                ("Eyebrow", "Rooms at Alpine Lodge"),
                ("Headline", "A Davis room for people who came to explore."),
                ("Body", "Private bathroom, refrigerator, flat-screen TV, Wi-Fi, DirecTV, and designated pet-friendly rooms give travelers the essentials between mountain days."),
                ("Primary CTA", "Reserve Now"),
                ("Secondary CTA", "Read FAQs"),
                ("Image alt text", "Guest room at Alpine Lodge in Davis, WV"),
            ],
        ),
        (
            "What To Expect",
            [
                ("Eyebrow", "What to expect"),
                ("Headline", "The room details travelers check first."),
                ("Body", "Alpine Lodge keeps the stay focused on the basics visitors use before and after waterfalls, ski days, trail time, and evenings in Davis or Thomas."),
                ("Amenities", "; ".join(data["amenities"])),
            ],
        ),
    ]
    for room in data["rooms"]:
        rooms_sections.append((f"Room Card: {room['title']}", [("Title", room["title"]), ("Caption", room["caption"]), ("Description", room["description"]), ("Image alt text", room["alt"])]))
    rooms_sections.append(
        (
            "Booking Note",
            [
                ("Eyebrow", "Direct booking"),
                ("Headline", "Check current rates and availability through Alpine Lodge reservations."),
                ("Body", "For pet-friendly rooms, group questions, early check-in requests, and current policies, reserve directly or contact the lodge before arrival."),
                ("CTA", "Check Availability"),
            ],
        )
    )
    add_page(doc, "Rooms Page", {"title": "Davis WV Hotel Rooms & Amenities", "description": "Book Davis, WV rooms at Alpine Lodge with private bathroom, refrigerator, flat-screen TV, Wi-Fi, DirecTV, and designated pet-friendly options."}, rooms_sections)
    append_md_page(md, "Rooms Page", {"title": "Davis WV Hotel Rooms & Amenities", "description": "Book Davis, WV rooms at Alpine Lodge with private bathroom, refrigerator, flat-screen TV, Wi-Fi, DirecTV, and designated pet-friendly options."}, rooms_sections)

    rates_sections = [
        (
            "Hero",
            [
                ("Eyebrow", "Rates and direct booking"),
                ("Headline", "Rooms from $99 on select dates, with live rates before you reserve."),
                ("Body", "Rates change by date, availability, room type, season, and local demand. Use the direct booking page for current Alpine Lodge rates."),
                ("Primary CTA", "Check Live Rates"),
                ("Secondary CTA", "View Rooms"),
                ("Image alt text", "Guest room at Alpine Lodge in Davis, West Virginia"),
            ],
        ),
        (
            "Rate Explanation",
            [
                ("Eyebrow", "Short answer"),
                ("Headline", "The booking engine is the current source for room price and availability."),
                ("Paragraph 1", "Alpine Lodge can be a fair-priced Davis basecamp, with rooms from $99 on select dates. The current rate depends on the dates, room type, availability, season, and local demand."),
                ("Paragraph 2", "For accurate pricing, guests should use the direct Cloudbeds reservation page rather than old search snippets or third-party summaries."),
            ],
        ),
        (
            "Rate Drivers",
            [
                ("Eyebrow", "What changes the rate"),
                ("Headline", "The same room can price differently depending on the trip."),
                ("Body", "These are the main factors guests should expect when comparing dates."),
                ("Factors", "Travel dates; Season; Room availability; Room type; Length of stay; Local event or ski weekend demand; Pet-friendly room availability"),
            ],
        ),
        (
            "Booking Note",
            [
                ("Eyebrow", "Direct reservations"),
                ("Headline", "Check the live room options before you commit."),
                ("Body", "The booking page shows current Alpine Lodge availability and rates for the selected dates. For pet-friendly rooms, group questions, or arrival questions, contact the lodge directly."),
                ("CTA", "Open Booking Site"),
            ],
        ),
    ]
    add_page(doc, "Rates Page", {"title": "Rates & Booking", "description": "Alpine Lodge rooms from $99 on select dates in Davis, WV. Check current room rates, availability, pet-friendly options, and booking details through direct Cloudbeds reservations."}, rates_sections)
    append_md_page(md, "Rates Page", {"title": "Rates & Booking", "description": "Alpine Lodge rooms from $99 on select dates in Davis, WV. Check current room rates, availability, pet-friendly options, and booking details through direct Cloudbeds reservations."}, rates_sections)

    crew_sections = [
        (
            "Hero",
            [
                ("Eyebrow", "Crew lodging in Davis, West Virginia"),
                ("Headline", "Crew rooms without the housing headache."),
                ("Body", "A practical Davis base for contractors, project teams, and traveling workers who need the essentials handled between workdays."),
                ("Primary CTA", "Check Live Availability"),
                ("Secondary CTA", "Contact The Lodge"),
                ("Image alt text", "Guest room at Alpine Lodge for crews working near Davis, West Virginia"),
            ],
        ),
        (
            "Positioning",
            [
                ("Eyebrow", "The short answer"),
                ("Headline", "Working near Davis, Thomas, Canaan Valley, or Tucker County?"),
                ("Paragraph 1", "Alpine Lodge gives crews a straightforward place to stay between workdays. Rooms include the practical basics workers use, and the lodge can discuss group booking options by phone."),
                ("Paragraph 2", "Public room rates start at $99 on select dates. Current rates, room types, and availability depend on the dates and should be confirmed through direct booking or with the lodge."),
            ],
        ),
        (
            "Useful Amenities",
            [
                ("Eyebrow", "Useful between shifts"),
                ("Headline", "The work-stay basics are already here."),
                ("Body", "A private room, a place to keep food and drinks cold, and simple property amenities for resetting after the workday."),
                ("Amenities", "Free Wi-Fi; In-room refrigerator; Private bathroom; Laundry machines; Ice machine; Irons and ironing boards; Flat-screen TV and DirecTV; Group booking options by phone"),
            ],
        ),
        (
            "Group Booking",
            [
                ("Eyebrow", "Booking several rooms?"),
                ("Headline", "Call with the dates, room count, and expected length of stay."),
                ("Body", "The lodge can confirm current availability and discuss group booking options. For a single room, the direct booking page is the fastest way to see live rates."),
                ("CTA 1", "Call The Lodge"),
                ("CTA 2", "Check Live Rates"),
            ],
        ),
        (
            "Business Traveler Review",
            [
                ("Eyebrow", "5/5 Google review from a business traveler"),
                ("Headline", "A work stay described as quiet and a great value."),
                ("Body", "The reviewer had planned a two-week stay for work before the job was cut short."),
                ("Review", "This place was wonderful."),
                ("Reviewer", "Bill Barowsky — 5/5 Google review"),
            ],
        ),
        (
            "Crew Stay FAQ",
            [
                ("Does Alpine Lodge accommodate contractors and work crews?", "Yes. Contractors, crews, and traveling workers can call Alpine Lodge to discuss room availability and group booking options for work in the Davis and Tucker County area."),
                ("What room amenities are useful for a longer work stay?", "Rooms include free Wi-Fi, a refrigerator, private bathroom, flat-screen TV, and DirecTV. The property also has laundry machines, an ice machine, and irons and ironing boards."),
                ("Are special group rates available for crews?", "Group booking options are available. Call Alpine Lodge with the work dates, number of rooms, and expected length of stay to discuss current options."),
                ("How do I check current rates and availability?", "Use the direct Cloudbeds booking page for live public rates and availability, or call Alpine Lodge to discuss a multi-room crew stay."),
            ],
        ),
    ]
    add_page(doc, "Crew Lodging Page", {"title": "Crew & Extended-Stay Lodging in Davis WV", "description": "Practical Davis, WV lodging for contractors, crews, and traveling workers, with Wi-Fi, refrigerators, laundry machines, ice, and group booking options."}, crew_sections)
    append_md_page(md, "Crew Lodging Page", {"title": "Crew & Extended-Stay Lodging in Davis WV", "description": "Practical Davis, WV lodging for contractors, crews, and traveling workers, with Wi-Fi, refrigerators, laundry machines, ice, and group booking options."}, crew_sections)

    catalog_sections = [
        (
            "Hero",
            [
                ("Eyebrow", "Agent-readable lodge knowledge"),
                ("Headline", "Alpine Lodge Knowledge Catalog"),
                ("Body", "A structured source of truth for rooms, rate guidance, amenities, nearby attractions, pet-friendly lodging, policies, FAQs, and Davis trip planning."),
                ("Primary CTA", "Reserve Now"),
                ("Secondary CTA", "View Rate Guidance"),
                ("Image alt text", "Alpine Lodge exterior in Davis, West Virginia"),
            ],
        ),
        (
            "Catalog Explanation",
            [
                ("Eyebrow", "What this catalog does"),
                ("Headline", "It helps people and AI systems understand the lodge without guessing."),
                ("Paragraph 1", "A practical, fair-priced Davis, WV basecamp for Blackwater Falls, Canaan Valley, Timberline Mountain, White Grass, Thomas, Dolly Sods, hiking, skiing, food, music, and low-cost mountain weekends."),
                ("Paragraph 2", "The catalog does not replace the website design. It gives the site a clearer structured layer for search, AI answer engines, trip planners, and future content updates."),
                ("Freshness note", f"Catalog last updated: {data['catalogUpdated']}."),
            ],
        ),
        (
            "Catalog Sections",
            [
                ("Eyebrow", "Catalog sections"),
                ("Headline", "The core lodging facts are now grouped into machine-readable categories."),
                ("Body", "These are the facts guests tend to ask for before they book: room basics, rate guidance, activities, policies, and current-source guardrails."),
            ] + [(item["name"], f"{item['summary']} CTA: Related page") for item in data["lodgingServices"]],
        ),
        (
            "Rate Guidance",
            [
                ("Eyebrow", "Rate guidance"),
                ("Headline", "Rates are explained clearly, but live pricing stays in the booking engine."),
                ("Body", "The catalog keeps the public price signal clear while preventing stale or invented rate details."),
            ] + [(item["name"], item["answer"]) for item in data["rateKnowledge"]],
        ),
        (
            "Policy Facts",
            [
                ("Eyebrow", "Policy facts"),
                ("Headline", "Guest policies are grouped where answer engines can find them."),
                ("Body", "This helps keep answers grounded in the site instead of letting outside tools infer details."),
            ] + [(item["name"], item["summary"]) for item in data["policyKnowledge"]],
        ),
        (
            "Machine-Readable Files",
            [
                ("Eyebrow", "Machine-readable files"),
                ("Headline", "Static files are available for crawlers and future workflows."),
                ("Body", "The catalog is also published as JSON and Markdown files so non-Google AI tools can parse it without relying on visual page rendering."),
                ("CTA 1", "Catalog JSON"),
                ("CTA 2", "Lodging Markdown"),
                ("CTA 3", "Check Live Rates"),
            ],
        ),
    ]
    add_page(doc, "Knowledge Catalog Page", {"title": "Knowledge Catalog", "description": "Structured Alpine Lodge lodging, rate, amenity, policy, FAQ, and Davis WV trip-planning knowledge for guests, search engines, and AI tools."}, catalog_sections)
    append_md_page(md, "Knowledge Catalog Page", {"title": "Knowledge Catalog", "description": "Structured Alpine Lodge lodging, rate, amenity, policy, FAQ, and Davis WV trip-planning knowledge for guests, search engines, and AI tools."}, catalog_sections)

    location_sections = [
        (
            "Hero",
            [
                ("Eyebrow", "Davis, WV location"),
                ("Headline", "Stay in Davis, then choose the day from there."),
                ("Body", "Alpine Lodge sits at 850 Sawmill Lane in Davis, West Virginia, close to waterfalls, ski areas, trails, food, music, and scenic drives."),
                ("Primary CTA", "Reserve Now"),
                ("Secondary CTA", "Plan Experiences"),
                ("Image alt text", "Downtown Davis, West Virginia near Alpine Lodge"),
            ],
        ),
        (
            "Address Panel",
            [
                ("Eyebrow", "Address"),
                ("Headline", "850 Sawmill Lane, Davis, WV 26260"),
                ("Body", "Use Alpine Lodge as the Davis center point for Blackwater Falls, Canaan Valley, Timberline Mountain, White Grass, Thomas, and downtown Davis."),
                ("CTA 1", "Open Map"),
                ("CTA 2", "Reserve Nearby"),
            ],
        ),
        (
            "Distance Rows",
            [(item["place"], f"{item['time']} — {item['note']}") for item in data["distances"]],
        ),
        (
            "Official Local Links",
            [
                ("Eyebrow", "Useful local links"),
                ("Headline", "Confirm hours, prices, tickets, and conditions with the official source."),
                ("Body", "Alpine Lodge helps frame the trip. Operators and official park resources provide the current day-of details."),
            ],
        ),
    ]
    add_page(doc, "Location Page", {"title": "Location", "description": "Stay at Alpine Lodge in Davis, WV, about 5 minutes from Blackwater Falls, 15 minutes from Canaan Valley, and 20 minutes from Timberline and White Grass."}, location_sections)
    append_md_page(md, "Location Page", {"title": "Location", "description": "Stay at Alpine Lodge in Davis, WV, about 5 minutes from Blackwater Falls, 15 minutes from Canaan Valley, and 20 minutes from Timberline and White Grass."}, location_sections)

    experiences_sections = [
        (
            "Hero",
            [
                ("Eyebrow", "Explore from Alpine Lodge"),
                ("Headline", "Things to do near Davis, WV when the room is not the whole trip."),
                ("Body", "Use Alpine Lodge as the Davis base for waterfalls, ski areas, trails, food, music, scenic drives, and pet-friendly mountain weekends."),
                ("Primary CTA", "Reserve Now"),
                ("Secondary CTA", "View Trip Guides"),
                ("Image alt text", "Blackwater Falls near Alpine Lodge in Davis, WV"),
            ],
        ),
        (
            "Experience Pages Intro",
            [
                ("Eyebrow", "Experience pages"),
                ("Headline", "Choose the trip first. Book the room around it."),
                ("Body", "Each page gives visitors the details they actually compare: drive time, likely cost level, best season, official links, and why Davis works as the base."),
            ],
        ),
        (
            "Official Planning Links",
            [
                ("Eyebrow", "Official planning links"),
                ("Headline", "Check changing details before you go."),
                ("Body", "Hours, prices, tickets, road access, trail conditions, and weather can change. Use these official sites for the current details."),
            ],
        ),
    ]
    add_page(doc, "Experiences Index Page", {"title": "Things To Do Near Davis WV", "description": "Explore Blackwater Falls, Canaan Valley, Timberline Mountain, White Grass, Dolly Sods, Thomas, hiking, skiing, pet-friendly trips, and low-cost Davis-area days from Alpine Lodge."}, experiences_sections)
    append_md_page(md, "Experiences Index Page", {"title": "Things To Do Near Davis WV", "description": "Explore Blackwater Falls, Canaan Valley, Timberline Mountain, White Grass, Dolly Sods, Thomas, hiking, skiing, pet-friendly trips, and low-cost Davis-area days from Alpine Lodge."}, experiences_sections)

    for item in data["experiences"]:
        sections = [
            (
                "Card / Listing Copy",
                [
                    ("Eyebrow", item["eyebrow"]),
                    ("Title", item["title"]),
                    ("Description", item["description"]),
                    ("Distance", item["distance"]),
                    ("Cost level", item["budget"]),
                    ("Card CTA", "Plan this trip"),
                    ("Compact card CTA", "View details"),
                    ("Image alt text", item["alt"]),
                ],
            ),
            (
                "Detail Page Hero",
                [
                    ("Eyebrow", item["eyebrow"]),
                    ("Headline", f"{item['title']} from Alpine Lodge"),
                    ("Body", item["description"]),
                    ("Primary CTA", "Reserve Now"),
                    ("Secondary CTA", "Ask The Planner"),
                ],
            ),
            ("Why It Works From Alpine Lodge", [(f"Paragraph {i + 1}", paragraph) for i, paragraph in enumerate(item["body"])]),
            ("Before You Go", [(f"Tip {i + 1}", tip) for i, tip in enumerate(item["tips"])]),
            (
                "Detail Sidebar",
                [
                    ("Distance", item["distance"]),
                    ("Cost level", item["budget"]),
                    ("Best season", item["season"]),
                    ("Good for", ", ".join(item["bestFor"])),
                    ("CTA 1", "Official Details"),
                    ("CTA 2", "Stay Nearby"),
                ],
            ),
            (
                "Related Section",
                [
                    ("Eyebrow", "Keep planning"),
                    ("Headline", "More ways to use Alpine Lodge as the Davis base."),
                ],
            ),
        ]
        add_page(doc, f"Experience Page: {item['title']}", {"title": item["seoTitle"], "description": item["description"]}, sections)
        append_md_page(md, f"Experience Page: {item['title']}", {"title": item["seoTitle"], "description": item["description"]}, sections)

    guide_index_sections = [
        (
            "Hero",
            [
                ("Eyebrow", "Practical trip guides"),
                ("Headline", "A better mountain weekend starts before you pack the car."),
                ("Body", "Use these guides to plan what to do, where to go, what to check, and how to use Alpine Lodge as the Davis room for the trip."),
                ("Primary CTA", "Reserve Now"),
                ("Secondary CTA", "Use Planner"),
                ("Image alt text", "Mountain scenery near Davis, West Virginia"),
            ],
        ),
        (
            "Guide Index Intro",
            [
                ("Eyebrow", "Choose a guide"),
                ("Headline", "Built for travelers who want the area more than the lobby."),
                ("Body", "These are not luxury itineraries. They are clear plans for people who want waterfalls, ski days, food, music, and a sensible room in Davis."),
                ("Card CTA", "Open guide"),
            ],
        ),
    ]
    add_page(doc, "Trip Guides Index Page", {"title": "Davis WV Trip Guides", "description": "Davis, WV trip guides for Alpine Lodge guests, including ski weekends, rainy days, pet-friendly travel, Blackwater Falls, Thomas, Canaan Valley, and local food."}, guide_index_sections)
    append_md_page(md, "Trip Guides Index Page", {"title": "Davis WV Trip Guides", "description": "Davis, WV trip guides for Alpine Lodge guests, including ski weekends, rainy days, pet-friendly travel, Blackwater Falls, Thomas, Canaan Valley, and local food."}, guide_index_sections)

    for guide in data["tripGuides"]:
        sections = [
            (
                "Card / Hero Copy",
                [
                    ("Eyebrow", "Trip guide"),
                    ("Title", guide["title"]),
                    ("Description", guide["description"]),
                    ("Best for", guide["audience"]),
                    ("Hero CTA", "Reserve Alpine Lodge"),
                    ("Card CTA", "Open guide"),
                ],
            )
        ]
        for day in guide["days"]:
            sections.append((day["label"], [(f"Step {i + 1}", item) for i, item in enumerate(day["items"])]))
        sections.append(
            (
                "Inline CTA",
                [
                    ("Headline", "Make Alpine Lodge the Davis room."),
                    ("Body", "Book the room in Davis, then use the rest of the budget for the trip itself."),
                    ("CTA", "Check Availability"),
                ],
            )
        )
        add_page(doc, f"Trip Guide Page: {guide['title']}", {"title": guide["seoTitle"], "description": guide["description"]}, sections)
        append_md_page(md, f"Trip Guide Page: {guide['title']}", {"title": guide["seoTitle"], "description": guide["description"]}, sections)

    trip_planner_sections = [
        (
            "Page Hero",
            [
                ("Eyebrow", "Alpine trip planner"),
                ("Headline", "Ask for a Davis-area trip plan."),
                ("Body", "Build a weekend around your dates, group, budget, pets, ski plans, waterfalls, food, music, or rainy weather."),
                ("CTA", "Reserve Your Basecamp"),
            ],
        ),
        (
            "Planner Interface",
            [
                ("Sidebar eyebrow", "Try a prompt"),
                ("Sidebar headline", "Build the day around your priorities."),
                ("Initial assistant message", "Tell me your dates, group size, interests, and budget. I will suggest a Davis-area plan with Alpine Lodge as the room for the trip."),
                ("Starter prompt 1", "Plan a two-night Davis weekend with waterfalls and good food."),
                ("Starter prompt 2", "We are coming for skiing and want to keep lodging costs down."),
                ("Starter prompt 3", "What can we do with a dog near Davis, WV?"),
                ("Starter prompt 4", "Give me a rainy-day plan around Davis and Thomas."),
                ("Input label", "Ask the Alpine Lodge trip planner"),
                ("Input placeholder", "Ask for a weekend plan, ski trip, dog-friendly route, or rainy-day ideas"),
                ("Loading message", "Planning your trip..."),
                ("Fallback message", "I could not reach the trip planner right now. Start with Blackwater Falls, downtown Davis, Thomas, and Canaan Valley, then reserve Alpine Lodge as the Davis room for the trip."),
                ("API fallback message", "AI planning is ready to connect once the client adds an OpenAI API key. For now: use Alpine Lodge as your Davis room, start with Blackwater Falls, add Canaan Valley or Timberline based on season, spend an evening in Davis or Thomas, and reserve directly through the Alpine Lodge booking link."),
                ("AI planner system instruction", "You are Alpine Lodge's Davis, WV trip planner. Recommend value-minded itineraries using only the provided source facts. Do not invent hours, prices, trail conditions, weather, pet policies, or availability. Link users to official operators for changing details. Always include Alpine Lodge as the Davis room for the trip and end with the direct booking URL."),
            ],
        ),
    ]
    add_page(doc, "Trip Planner Page", {"title": "Davis WV Trip Planner", "description": "Plan a Davis, WV trip with Alpine Lodge as the room for Blackwater Falls, Canaan Valley, Timberline Mountain, Thomas, hiking, skiing, food, and music."}, trip_planner_sections)
    append_md_page(md, "Trip Planner Page", {"title": "Davis WV Trip Planner", "description": "Plan a Davis, WV trip with Alpine Lodge as the room for Blackwater Falls, Canaan Valley, Timberline Mountain, Thomas, hiking, skiing, food, and music."}, trip_planner_sections)

    faq_sections = [
        (
            "Hero",
            [
                ("Eyebrow", "Alpine Lodge FAQ"),
                ("Headline", "Answers before you book the Davis room."),
                ("Body", "Use these answers to plan the room, the drive, the pets, the ski weekend, and the local activities."),
                ("Primary CTA", "Reserve Now"),
                ("Secondary CTA", "Contact The Lodge"),
                ("Image alt text", "Alpine Lodge office in Davis, West Virginia"),
            ],
        ),
        ("FAQ Items", [(faq["question"], faq["answer"]) for faq in data["faqs"]]),
    ]
    add_page(doc, "FAQ Page", {"title": "Davis WV Hotel FAQ: Pets, Wi-Fi & Policies", "description": "Answers about booking Alpine Lodge in Davis, WV, including pets, Wi-Fi, room amenities, check-in, ski access, Blackwater Falls, group bookings, and local planning."}, faq_sections)
    append_md_page(md, "FAQ Page", {"title": "Davis WV Hotel FAQ: Pets, Wi-Fi & Policies", "description": "Answers about booking Alpine Lodge in Davis, WV, including pets, Wi-Fi, room amenities, check-in, ski access, Blackwater Falls, group bookings, and local planning."}, faq_sections)

    contact_sections = [
        (
            "Hero",
            [
                ("Eyebrow", "Contact Alpine Lodge"),
                ("Headline", "Book the room. Ask the question. Start from Davis."),
                ("Body", "Use direct reservations for availability, or contact Alpine Lodge for room, pet, group, and arrival questions."),
                ("Primary CTA", "Reserve Now"),
                ("Secondary CTA", "View Location"),
                ("Image alt text", "Exterior of Alpine Lodge in Davis, West Virginia"),
            ],
        ),
        (
            "Contact Cards",
            [
                ("Address", "850 Sawmill Lane, Davis, WV 26260"),
                ("Phone", site["phone"]),
                ("Email", site["email"]),
            ],
        ),
        (
            "Booking Note",
            [
                ("Eyebrow", "Fastest path"),
                ("Headline", "For rates and availability, use the current Alpine Lodge booking site."),
                ("Body", "The Reserve Now button opens the active Cloudbeds reservation page in a new tab."),
                ("CTA", "Open Booking Site"),
            ],
        ),
    ]
    add_page(doc, "Contact Page", {"title": "Contact & Directions", "description": "Contact Alpine Lodge in Davis, WV. Find the address, phone, email, direct booking link, and directions for lodging near Blackwater Falls and Canaan Valley."}, contact_sections)
    append_md_page(md, "Contact Page", {"title": "Contact & Directions", "description": "Contact Alpine Lodge in Davis, WV. Find the address, phone, email, direct booking link, and directions for lodging near Blackwater Falls and Canaan Valley."}, contact_sections)

    photo_credit_sections = [
        (
            "Hero",
            [
                ("Eyebrow", "Image sources"),
                ("Headline", "Photo credits"),
                ("Body", "Destination images help visitors see the real places around Alpine Lodge. External photographs are credited here with their original source and license where available."),
                ("Primary CTA", "Reserve Now"),
                ("Secondary CTA", "View Experiences"),
                ("Image alt text", "Misty autumn road in Dolly Sods Wilderness"),
            ],
        ),
        (
            "Credits",
            [
                ("Dolly Sods autumn road", "Photo by Bold Frontiers. Creative Commons Attribution 3.0 Unported."),
                ("East Avenue in Thomas, West Virginia", "Photo by Tim Kiser. Creative Commons Attribution-ShareAlike 2.5."),
                ("White Grass Ski Touring Center", "Photo by White Grass Ski Touring Center. Used from the official experience provider website."),
                ("Link labels", "Original Source; license name"),
            ],
        ),
    ]
    add_page(doc, "Photo Credits Page", {"title": "Photo Credits", "description": "Photo sources and Creative Commons attribution for destination images used on the Alpine Lodge website."}, photo_credit_sections)
    append_md_page(md, "Photo Credits Page", {"title": "Photo Credits", "description": "Photo sources and Creative Commons attribution for destination images used on the Alpine Lodge website."}, photo_credit_sections)

    not_found_sections = [
        (
            "404 Page",
            [
                ("Eyebrow", "Page not found"),
                ("Headline", "This trail does not go where we thought."),
                ("Body", "Head back to Alpine Lodge planning pages or reserve a room for your Davis trip."),
                ("CTA 1", "Back Home"),
                ("CTA 2", "Reserve Now"),
            ],
        )
    ]
    add_page(doc, "Not Found Page", None, not_found_sections)
    append_md_page(md, "Not Found Page", None, not_found_sections)

    partner_sections = [
        (
            "Local Partner Links",
            [(item["name"], f"{item['type']} — {item['href']}") for item in data["partnerLinks"]],
        ),
        (
            "Social Links",
            [(item["label"], item["href"]) for item in site["socials"]],
        ),
        (
            "Direct Booking",
            [("Cloudbeds booking URL", site["bookingUrl"])],
        ),
    ]
    add_page(doc, "Reference Links", None, partner_sections)
    append_md_page(md, "Reference Links", None, partner_sections)

    ai_text_files = ["llms.txt", "lodging.md", "rates.md", "policies.md", "faq.md", "work-stays.md"]
    ai_text_sections = []
    for filename in ai_text_files:
        source_path = ROOT / "public" / filename
        if source_path.exists():
            ai_text_sections.append(
                (f"Public File: /{filename}", [("Copy", source_path.read_text(encoding="utf-8"))])
            )
    if ai_text_sections:
        add_page(doc, "AI Search Copy", None, ai_text_sections)
        append_md_page(md, "AI Search Copy", None, ai_text_sections)

    doc.save(OUTPUT_DOCX)
    OUTPUT_MD.write_text("".join(md), encoding="utf-8")


if __name__ == "__main__":
    build_document(load_site_data())
    print(OUTPUT_DOCX)
    print(OUTPUT_MD)
